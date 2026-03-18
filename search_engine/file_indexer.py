""" 
File Indexer 
Scans folders and extracts content from PDF, DOCX, TXT, and EXE files using parallel processing. 
"""

import os
import hashlib
import re
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
import fitz

# Suppress annoying PDFMiner warnings 
logging.getLogger("pdfminer").setLevel(logging.ERROR)


class FileIndexer:
    def __init__(self):
        """Initialize the file indexer."""
        # 1. ADDED .exe HERE 
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.exe'}

        self.system_folders = {
            'C:\\WINDOWS',
            'C:\\PROGRAM FILES',
            'C:\\PROGRAM FILES (X86)',
            'C:\\SYSTEM32',
            'C:\\PROGRAMDATA',
            'C:\\USERS\\DEFAULT',
            'C:\\BOOT',
            'C:\\RECOVERY'
        }

    def scan_directory(self, folder_path):
        """Scans the folder and returns a list of supported files."""
        if not os.path.exists(folder_path):
            raise ValueError(f"Folder does not exist: {folder_path}")

        all_files = []
        for root, dirs, files in os.walk(folder_path):
            # Filter directories in-place 
            dirs[:] = [d for d in dirs if not self._should_skip_folder(os.path.join(root, d))]

            for file in files:
                _, ext = os.path.splitext(file.lower())
                if ext in self.supported_extensions:
                    all_files.append(os.path.join(root, file))

        print(f"Found {len(all_files)} supported files to scan.")
        return all_files

    def process_files(self, file_paths, existing_ids=None):
        """Process a list of files in parallel and yield documents."""
        if existing_ids is None:
            existing_ids = set()

        print(f"Processing {len(file_paths)} files with {len(existing_ids)} existing IDs...")

        files_to_process = []
        skipped_count = 0

        for f in file_paths:
            try:
                # Pre-calculate ID to check if we can skip 
                file_stats = os.stat(f)
                doc_id = hashlib.md5(f"{f}_{file_stats.st_mtime}".encode()).hexdigest()

                if doc_id in existing_ids:
                    skipped_count += 1
                    continue

                files_to_process.append(f)
            except:
                continue

        print(f"Skipping {skipped_count} already indexed files. Processing {len(files_to_process)} new/modified files.")

        # ThreadPoolExecutor is best for I/O bound tasks like file reading 
        max_workers = min(32, (os.cpu_count() or 1) * 4)
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {executor.submit(self._process_single_path_independent, f): f for f in files_to_process}

            for future in as_completed(future_to_file):
                try:
                    result = future.result()
                    if result:
                        yield result
                except Exception:
                    pass

    def _process_single_path_independent(self, file_path):
        """Process a single file: extract text and metadata."""
        try:
            _, ext = os.path.splitext(file_path.lower())

            content = self._extract_content(file_path, ext)

            if not content or not content.strip():
                # If content is empty, use filename as content so it's still searchable 
                content = os.path.basename(file_path)

            file_stats = os.stat(file_path)
            doc_id = hashlib.md5(f"{file_path}_{file_stats.st_mtime}".encode()).hexdigest()

            return {
                "id": doc_id,
                "content": content,
                "metadata": {
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "modified": file_stats.st_mtime,
                    "size": file_stats.st_size,
                    "type": ext
                }
            }
        except Exception:
            return None

    def _extract_content(self, file_path, extension):
        """Extract text content from a file based on its extension."""
        if extension == '.txt':
            return self._extract_txt_content(file_path)
        elif extension == '.pdf':
            return self._extract_pdf_content(file_path)
        elif extension == '.docx':
            return self._extract_docx_content(file_path)
        elif extension == '.exe':
            return self._extract_exe_content(file_path)  # 2. ADDED HANDLER HERE 
        return ""

    # --- NEW: EXE EXTRACTOR --- 
    def _extract_exe_content(self, file_path):
        """ 
        Extracts readable ASCII and Unicode strings from binary files. 
        This allows searching for metadata inside EXEs without running them. 
        """
        try:
            with open(file_path, 'rb') as f:
                # Read only the first 2MB to ensure speed
                data = f.read(2 * 1024 * 1024)

                # Regex to find sequences of 4+ printable characters 
                strings = re.findall(b'[ -~]{4,}', data)

                # Decode and join 
                text = " ".join([s.decode('utf-8', 'ignore') for s in strings])

                # Limit content size to avoid polluting the database with garbage 
                return self._clean_text(text[:5000])
        except:
            return ""

    def _extract_txt_content(self, file_path):
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return self._clean_text(file.read())
            except:
                continue
        return ""

    def _extract_pdf_content(self, file_path):
        """ 
        Extracts text from PDF using PyMuPDF (Fitz). 
        This is 20x faster than PDFMiner. 
        """
        try:
            text_content = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_content.append(page.get_text())

            return self._clean_text("\n".join(text_content))
        except Exception:
            return ""

    def _extract_docx_content(self, file_path):
        try:
            doc = Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return self._clean_text('\n'.join(full_text))
        except:
            return ""

    def _clean_text(self, text):
        if not text:
            return ""
        # Remove excessive whitespace 
        text = re.sub(r'\s+', ' ', text)
        # Remove null bytes and non-printable control chars 
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        return text.strip()[:100000]

    def _should_skip_folder(self, folder_path):
        if not folder_path:
            return False

        name = os.path.basename(folder_path)

        if name.startswith('.'):
            return True

        ignore_names = {
            'node_modules', 'site-packages', 'dist-info', '__pycache__',
            'venv', 'env', 'odf_env', 'libs', 'include', 'scripts', 'bin', 'obj'
        }
        if name.lower() in ignore_names:
            return True

        norm_path = os.path.normpath(folder_path.upper())
        for sys_folder in self.system_folders:
            if norm_path.startswith(sys_folder):
                return True

        return False
